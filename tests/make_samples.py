# -*- coding: utf-8 -*-
"""生成各格式测试样本（原创内容），并跑核心模块冒烟测试。"""
import os
import sys
import tempfile
import zipfile

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, BASE)

SAMPLE_DIR = os.path.join(BASE, "sample")

# ---------- 原创样章 ----------
def make_txt_chapters():
    chs = []
    for n in range(1, 7):
        title = f"第{n}章 星火"
        body = (
            "\n\n".join(
                [
                    "远处的城市在暮色里亮起第一盏灯。林远把背包放在长椅上，望着天边那道若隐若现的星轨。",
                    "他记得爷爷说过，每一颗星星都是一艘远航的船，载着某个再也回不去的人。",
                    "风从东边吹来，带来潮湿的泥土气息。林远深吸一口气，朝那座亮着灯的楼走去。",
                    "门开了，一个白发老人站在阴影里，手里捏着一枚发光的硬币。老人说，你终于来了。",
                    "林远没有回答，只是把那枚硬币接过来，硬币在他掌心转了个圈，化作一团蓝色的火焰。",
                ]
            )
            + "\n"
        )
        chs.append((title, body))
    return chs


def write_txt():
    parts = []
    for t, b in make_txt_chapters():
        parts.append(t + "\n" + b)
    text = "".join(parts)
    p = os.path.join(SAMPLE_DIR, "星火.txt")
    with open(p, "w", encoding="utf-8") as f:
        f.write(text)
    return p


def write_epub():
    from html import escape

    p = os.path.join(SAMPLE_DIR, "星火.epub")
    chapters = make_txt_chapters()
    with zipfile.ZipFile(p, "w") as z:
        z.writestr("mimetype", "application/epub+zip")
        z.writestr(
            "META-INF/container.xml",
            '<?xml version="1.0"?><container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
            '<rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles></container>',
        )
        manifest = ['<item id="ch0" href="ch0.xhtml" media-type="application/xhtml+xml"/>']
        spine = ['<itemref idref="ch0"/>']
        files = []
        for i, (t, b) in enumerate(chapters):
            manifest.append(f'<item id="ch{i+1}" href="ch{i+1}.xhtml" media-type="application/xhtml+xml"/>')
            spine.append(f'<itemref idref="ch{i+1}"/>')
            paras = "".join(f"<p>{escape(x)}</p>" for x in b.strip().split("\n\n"))
            files.append(
                (f"ch{i+1}.xhtml", f'<html xmlns="http://www.w3.org/1999/xhtml"><head><title>{escape(t)}</title></head><body><h1>{escape(t)}</h1>{paras}</body></html>')
            )
        opf = (
            '<?xml version="1.0" encoding="utf-8"?>'
            '<package xmlns="http://www.idpf.org/2007/opf" version="2.0" unique-identifier="uid">'
            '<metadata xmlns:dc="http://purl.org/dc/elements/1.1/">'
            '<dc:title>星火</dc:title><dc:creator>测试作者</dc:creator>'
            '<dc:identifier id="uid">sample-001</dc:identifier></metadata>'
            '<manifest>'
            + "".join(manifest)
            + '</manifest><spine toc="ncx">'
            + "".join(spine)
            + "</spine></package>"
        )
        z.writestr("OEBPS/content.opf", opf)
        for name, content in files:
            z.writestr(f"OEBPS/{name}", content)
    return p


def write_docx():
    import docx

    p = os.path.join(SAMPLE_DIR, "星火.docx")
    d = docx.Document()
    d.core_properties.title = "星火"
    d.core_properties.author = "测试作者"
    for t, b in make_txt_chapters():
        d.add_heading(t, level=1)
        for para in b.strip().split("\n\n"):
            d.add_paragraph(para)
    d.save(p)
    return p


def write_pdf():
    import pymupdf as fitz

    p = os.path.join(SAMPLE_DIR, "星火.pdf")
    doc = fitz.open()
    page = doc.new_page()
    y = 60
    for t, b in make_txt_chapters():
        page.insert_text((60, y), t, fontsize=16)
        y += 30
        for para in b.strip().split("\n\n"):
            page.insert_text((60, y), para[:60], fontsize=11)
            y += 18
        y += 12
    doc.save(p)
    doc.close()
    return p


def main():
    os.makedirs(SAMPLE_DIR, exist_ok=True)
    paths = {
        "txt": write_txt(),
        "epub": write_epub(),
        "docx": write_docx(),
        "pdf": write_pdf(),
    }
    for k, p in paths.items():
        print(f"[sample] {k}: {os.path.getsize(p)} bytes -> {p}")
    return paths


if __name__ == "__main__":
    main()
